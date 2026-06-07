from __future__ import annotations

import plistlib
import re
import subprocess
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path("/Users/td/VscodeProjects/ai_project/task_spiter/new-project")
SOURCE_DIR = Path("/Users/td/Documents/ppt整理")
TMP_DIR = ROOT / "tmp_exam_pack"
OUTPUT_DIR = ROOT / "output"
OUTPUT_DOCX = OUTPUT_DIR / "公文写作与处理-开卷考试整理版.docx"


FILES = [
    "事务类公文写作：计划与总结.pptx",
    "决策指辉类.pptx",
    "第一章 公文概述.ppt",
    "第四章 报告请示 第六章 批复.ppt",
    "第三章公文写作概述.ppt",
    "第六章 通知 第七章 通报.ppt",
    "第六章  意见.ppt",
    "第二章公文格式.ppt",
    "第八章 商洽、纪要类公文.ppt",
]

REFERENCE_PDF = "《公文写作与处理（卫生行政执法文书部分）》期末考试参考课件.pdf"

KNOWLEDGE_HINTS = (
    "定义",
    "特点",
    "作用",
    "分类",
    "要求",
    "结构",
    "格式",
    "写法",
    "写作方法",
    "注意事项",
    "适用范围",
    "行文方向",
    "基本要求",
    "概念",
    "原则",
    "类型",
    "文种",
    "标题",
    "正文",
    "结尾",
    "开头",
    "语言",
    "写作要求",
    "基本格式",
)

CASE_HINTS = (
    "案例",
    "例",
    "病文",
    "改写",
    "分析",
    "思考",
    "练习",
    "判断",
    "多选",
    "单选",
    "材料",
    "示例",
    "范文",
    "启示",
    "实训",
    "实战",
    "问题",
    "错误",
)


@dataclass
class SourceText:
    name: str
    text: str


def run(cmd: list[str]) -> subprocess.CompletedProcess[str]:
    return subprocess.run(cmd, check=True, text=True, capture_output=True)


def mdimport_text(path: Path) -> str:
    out_file = TMP_DIR / f"{path.name}.plist"
    subprocess.run(
        ["mdimport", "-t", "-o", str(out_file), str(path)],
        check=True,
        text=True,
        capture_output=True,
    )
    raw = out_file.read_bytes()
    text = raw.decode("utf-8", errors="ignore")
    match = re.search(r'kMDItemTextContent = "(.*)";\s*}', text, re.S)
    if not match:
        raise RuntimeError(f"未能从 {path} 提取 kMDItemTextContent")
    value = match.group(1)
    # Spotlight 输出里混合了 \UXXXX 与普通字符，需定向解码。
    value = value.replace("\\n", "\n")
    value = value.replace('\\"', '"')
    value = decode_spotlight_escapes(value)
    return clean_text(value)


def decode_spotlight_escapes(value: str) -> str:
    def repl(match: re.Match[str]) -> str:
        code = match.group(1)
        try:
            return chr(int(code, 16))
        except ValueError:
            return match.group(0)

    value = re.sub(r"\\U([0-9A-Fa-f]{4})", repl, value)
    value = value.replace("\\t", "\t")
    return value


def pdf_text(path: Path) -> str:
    out_file = TMP_DIR / f"{path.stem}.txt"
    subprocess.run(
        ["pdftotext", "-layout", str(path), str(out_file)],
        check=True,
        text=True,
        capture_output=True,
    )
    return clean_text(out_file.read_text(encoding="utf-8", errors="ignore"))


def clean_text(text: str) -> str:
    text = text.replace("\r", "\n")
    text = text.replace("\x0c", "\n")
    text = text.replace("\u00a0", " ")
    text = text.replace("•", "\n• ")
    text = text.replace("", "\n■ ")
    text = text.replace("", "\n• ")
    text = text.replace("①", "\n1. ")
    text = text.replace("②", "\n2. ")
    text = text.replace("③", "\n3. ")
    text = text.replace("④", "\n4. ")
    text = text.replace("⑤", "\n5. ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def split_lines(text: str) -> list[str]:
    lines = []
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        if re.fullmatch(r"\d{1,3}", line):
            continue
        if re.fullmatch(r"20\d{2}/\d{1,2}/\d{1,2}", line):
            continue
        lines.append(line)
    return lines


def classify_lines(lines: Iterable[str]) -> tuple[list[str], list[str]]:
    knowledge: list[str] = []
    cases: list[str] = []
    for line in lines:
        normalized = line.replace(" ", "")
        if any(h in normalized for h in CASE_HINTS):
            cases.append(line)
        elif any(h in normalized for h in KNOWLEDGE_HINTS):
            knowledge.append(line)
        elif re.search(r"[A-DＡ-Ｄ][.、]", line) or "判断：" in line or "多选" in line:
            cases.append(line)
        else:
            knowledge.append(line)
    return knowledge, cases


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def set_default_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Songti SC"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
    normal.font.size = Pt(10.5)
    for section in doc.sections:
        section.top_margin = Pt(50)
        section.bottom_margin = Pt(50)
        section.left_margin = Pt(52)
        section.right_margin = Pt(52)


def add_bullets(doc: Document, items: list[str], limit: int | None = None) -> None:
    count = 0
    for item in items:
        if limit is not None and count >= limit:
            break
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
        count += 1


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def build_tip_section(doc: Document) -> None:
    add_heading(doc, "考试速查与做题技巧", level=1)
    intro = doc.add_paragraph()
    intro.add_run(
        "以下技巧是根据课件中的文种特征、格式要求和常见题型整理，便于开卷考试时快速定位。"
    )

    add_heading(doc, "一、选择题与判断题", level=2)
    add_bullets(
        doc,
        [
            "先看题干问的是定义、特点、适用范围、行文方向、结构还是格式，公文题最常考这些稳定知识点。",
            "凡是涉及“请示”“报告”“批复”“通知”“通报”“意见”“函”“纪要”的，先判断它们属于上行文、平行文还是下行文。",
            "判断题特别注意绝对化表述，如“必须只能一文一事”“任何情况下都能平行行文”等，往往是命题点。",
            "看到格式题，优先回忆党政机关公文格式要素：标题、主送机关、正文、成文日期、印章、附件说明、附注等。",
        ],
    )

    add_heading(doc, "二、公文标题拟写", level=2)
    add_bullets(
        doc,
        [
            "标题通常由发文机关、事由、文种构成，考试中最容易丢分的是文种误用或事由不准。",
            "先判定事项性质，再选文种：汇报工作看是否用“报告”，请求指示或批准通常用“请示”，答复请示用“批复”，周知事项常见“通知”或“通报”。",
            "标题要庄重、准确、简明，避免口语化、文学化、情绪化表达。",
        ],
    )

    add_heading(doc, "三、病文分析与改写", level=2)
    add_bullets(
        doc,
        [
            "先找硬伤：文种错误、行文方向错误、标题不规范、结构缺项、主送机关不当、语气不符。",
            "再找语言问题：空话套话过多、逻辑不清、层次混乱、表述含糊、错别字与标点不规范。",
            "改写时先保文种正确，再补结构完整，最后优化语言。",
        ],
    )

    add_heading(doc, "四、材料写作", level=2)
    add_bullets(
        doc,
        [
            "先审题定文种，再列提纲。不要一上来就写正文。",
            "材料题常考通知、请示、报告、通报、总结、计划等，要先判断写给谁、为了解决什么问题、希望对方做什么。",
            "正文尽量做到“背景/依据 - 事项/问题 - 措施/请求 - 结语”清楚完整。",
            "计划重在面向未来、目标措施清楚；总结重在回顾工作、提炼成绩经验、分析问题和今后打算。",
        ],
    )


def build_document(sources: list[SourceText], reference: str) -> None:
    doc = Document()
    set_default_style(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run("公文写作与处理开卷考试整理版")
    run_title.bold = True
    run_title.font.size = Pt(18)
    run_title.font.name = "Songti SC"
    run_title._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("依据课堂 PPT 全量整理，适合纸质打印翻阅").italic = True

    preface = doc.add_paragraph()
    preface.add_run(
        "说明：本资料以课件原文为核心，删除了教师姓名等与考试无关的信息；为方便开卷考试，前半部分重组为知识点与案例速查，后半部分保留各课件的完整文本整理稿，尽量保证不遗漏。"
    )

    all_knowledge: list[str] = []
    all_cases: list[str] = []

    per_source: dict[str, tuple[list[str], list[str], list[str]]] = {}
    for source in sources:
        lines = split_lines(source.text)
        knowledge, cases = classify_lines(lines)
        per_source[source.name] = (lines, knowledge, cases)
        all_knowledge.extend(knowledge)
        all_cases.extend(cases)

    add_heading(doc, "第一部分 知识点总整理", level=1)
    p = doc.add_paragraph()
    p.add_run("这一部分按课件分组浓缩主要概念、特点、分类、格式和写作要求，便于考试时快速查找。")

    for source in sources:
        lines, knowledge, _ = per_source[source.name]
        add_heading(doc, source.name, level=2)
        seen = set()
        cleaned: list[str] = []
        for item in knowledge:
            key = re.sub(r"\s+", "", item)
            if len(key) < 4 or key in seen:
                continue
            seen.add(key)
            cleaned.append(item)
        add_bullets(doc, cleaned, limit=120)

    add_heading(doc, "第二部分 案例、练习、病文与示例", level=1)
    p = doc.add_paragraph()
    p.add_run("这一部分把课件中带有案例、练习、病文、判断、多选、示例等内容单独汇总，方便考前翻案例。")
    for source in sources:
        _, _, cases = per_source[source.name]
        if not cases:
            continue
        add_heading(doc, source.name, level=2)
        seen = set()
        cleaned = []
        for item in cases:
            key = re.sub(r"\s+", "", item)
            if len(key) < 3 or key in seen:
                continue
            seen.add(key)
            cleaned.append(item)
        add_numbered(doc, cleaned[:100])

    build_tip_section(doc)

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "附录 各课件完整文本整理", level=1)
    p = doc.add_paragraph()
    p.add_run("以下为各 PPT/PPTX 及参考 PDF 提取后的完整文本整理稿，便于需要时进行全文检索式翻阅。")

    for source in sources:
        add_heading(doc, source.name, level=2)
        for para in split_full_text(source.text):
            doc.add_paragraph(para)

    add_heading(doc, "参考课件 PDF（交叉核对材料）", level=2)
    for para in split_full_text(reference):
        doc.add_paragraph(para)

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_DOCX))


def split_full_text(text: str) -> list[str]:
    chunks: list[str] = []
    current: list[str] = []
    for line in text.splitlines():
        line = line.strip()
        if not line:
            if current:
                chunks.append(" ".join(current))
                current = []
            continue
        if len(" ".join(current + [line])) > 120:
            if current:
                chunks.append(" ".join(current))
            current = [line]
        else:
            current.append(line)
    if current:
        chunks.append(" ".join(current))
    return chunks


def main() -> None:
    TMP_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    sources: list[SourceText] = []
    for file_name in FILES:
        path = SOURCE_DIR / file_name
        text = mdimport_text(path)
        (TMP_DIR / f"{file_name}.txt").write_text(text, encoding="utf-8")
        sources.append(SourceText(file_name, text))

    reference_text = pdf_text(SOURCE_DIR / REFERENCE_PDF)
    (TMP_DIR / f"{REFERENCE_PDF}.txt").write_text(reference_text, encoding="utf-8")

    build_document(sources, reference_text)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
